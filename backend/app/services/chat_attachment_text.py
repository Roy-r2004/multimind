"""Extract bounded text excerpts from chat attachment files."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from app.core.exceptions import InvalidAttachmentError
from app.core.logging import get_logger

logger = get_logger(__name__)

ATTACHMENT_TEXT_EXCERPT_MAX = 20_000
_XLSX_MAX_SHEETS = 5
_XLSX_MAX_ROWS = 200
_XLSX_MAX_COLS = 30
_PDF_MAX_PAGES = 100
_TRUNCATION_MARKER = "\n[Content truncated]"
_BLANK_LINE_RE = re.compile(r"\n{3,}")


def extract_attachment_text(content: bytes, extension: str) -> tuple[str | None, str]:
    """Return (excerpt, excerpt_status) for an allowed attachment extension."""
    ext = extension.lower()
    try:
        if ext in {".docx"}:
            return _finalize(_extract_docx_text(content))
        if ext in {".xlsx"}:
            return _finalize(_extract_xlsx_text(content))
        if ext in {".pdf"}:
            return _finalize(_extract_pdf_text(content))
        return _finalize(_extract_plain_text(content))
    except InvalidAttachmentError:
        logger.warning(
            "chat_attachment_extraction_failed",
            extension=ext,
        )
        raise


def extract_attachment_text_from_path(path: Path, extension: str) -> tuple[str | None, str]:
    """Extract from a temp/final file path without requiring a full in-memory upload buffer.

    Plain text reads only a bounded prefix. Office/PDF parsers open the path directly.
    """
    ext = extension.lower()
    try:
        if ext in {".docx"}:
            return _finalize(_extract_docx_from_path(path))
        if ext in {".xlsx"}:
            return _finalize(_extract_xlsx_from_path(path))
        if ext in {".pdf"}:
            return _finalize(_extract_pdf_from_path(path))
        return _finalize(_extract_plain_text_from_path(path))
    except InvalidAttachmentError:
        logger.warning(
            "chat_attachment_extraction_failed",
            extension=ext,
        )
        raise


def _finalize(text: str) -> tuple[str | None, str]:
    excerpt = text[:ATTACHMENT_TEXT_EXCERPT_MAX]
    if not excerpt.strip():
        return None, "empty"
    return excerpt, "ready"


def excerpt_from_transcript(text: str) -> tuple[str | None, str]:
    """Map a transcription string into attachment text_excerpt / excerpt_status."""
    return _finalize(text.strip())


def _extract_plain_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _extract_plain_text_from_path(path: Path) -> str:
    # Bound the read so large text files do not need a full in-memory copy.
    with path.open("rb") as handle:
        raw = handle.read(ATTACHMENT_TEXT_EXCERPT_MAX + 4096)
    return raw.decode("utf-8", errors="replace")


def _extract_docx_text(content: bytes) -> str:
    return _docx_document_to_text(_open_docx(BytesIO(content)))


def _extract_docx_from_path(path: Path) -> str:
    return _docx_document_to_text(_open_docx(str(path)))


def _open_docx(source: BytesIO | str):
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as exc:  # pragma: no cover - dependency should be installed
        raise InvalidAttachmentError("DOCX support is unavailable on this server") from exc

    try:
        return Document(source)
    except PackageNotFoundError as exc:
        raise InvalidAttachmentError("Invalid or corrupt DOCX file") from exc
    except Exception as exc:
        raise InvalidAttachmentError("Invalid or corrupt DOCX file") from exc


def _docx_document_to_text(document) -> str:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(parts)


def _extract_xlsx_text(content: bytes) -> str:
    return _xlsx_workbook_to_text(_open_xlsx(BytesIO(content)))


def _extract_xlsx_from_path(path: Path) -> str:
    return _xlsx_workbook_to_text(_open_xlsx(str(path)))


def _open_xlsx(source: BytesIO | str):
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise InvalidAttachmentError("XLSX support is unavailable on this server") from exc

    try:
        return load_workbook(
            source,
            read_only=True,
            data_only=True,
        )
    except Exception as exc:
        message = str(exc).lower()
        if "password" in message or "encrypted" in message:
            raise InvalidAttachmentError(
                "Password-protected or encrypted Excel files are not supported"
            ) from exc
        raise InvalidAttachmentError("Invalid or corrupt XLSX file") from exc


def _xlsx_workbook_to_text(workbook) -> str:
    parts: list[str] = []
    truncated = False
    try:
        sheet_names = list(workbook.sheetnames)[:_XLSX_MAX_SHEETS]
        if len(workbook.sheetnames) > _XLSX_MAX_SHEETS:
            truncated = True

        for sheet_name in sheet_names:
            worksheet = workbook[sheet_name]
            parts.append(f"Worksheet: {sheet_name}")
            row_count = 0
            for row in worksheet.iter_rows(max_col=_XLSX_MAX_COLS, values_only=True):
                values = [_cell_to_text(cell) for cell in row]
                if not any(value.strip() for value in values):
                    continue
                row_count += 1
                if row_count > _XLSX_MAX_ROWS:
                    truncated = True
                    break
                parts.append("\t".join(values).rstrip("\t"))
                joined = "\n".join(parts)
                if len(joined) >= ATTACHMENT_TEXT_EXCERPT_MAX:
                    truncated = True
                    break
            if truncated and len("\n".join(parts)) >= ATTACHMENT_TEXT_EXCERPT_MAX:
                break
    finally:
        workbook.close()

    text = "\n".join(parts).strip()
    if truncated:
        budget = ATTACHMENT_TEXT_EXCERPT_MAX - len(_TRUNCATION_MARKER)
        if budget < 0:
            return _TRUNCATION_MARKER.strip()
        text = text[:budget].rstrip() + _TRUNCATION_MARKER
    return text


def _cell_to_text(value: object) -> str:
    if value is None:
        return ""
    return str(value).replace("\t", " ").replace("\n", " ").strip()


def _extract_pdf_text(content: bytes) -> str:
    return _pdf_reader_to_text(_open_pdf(BytesIO(content)))


def _extract_pdf_from_path(path: Path) -> str:
    return _pdf_reader_to_text(_open_pdf(str(path)))


def _open_pdf(source: BytesIO | str):
    try:
        from pypdf import PdfReader
        from pypdf.errors import FileNotDecryptedError, PdfReadError
    except ImportError as exc:  # pragma: no cover
        raise InvalidAttachmentError("PDF support is unavailable on this server") from exc

    try:
        reader = PdfReader(source)
    except PdfReadError as exc:
        raise InvalidAttachmentError("Invalid or corrupt PDF file") from exc
    except Exception as exc:
        message = str(exc).lower()
        if "password" in message or "encrypt" in message or "decrypt" in message:
            raise InvalidAttachmentError(
                "Password-protected or encrypted PDFs are not supported"
            ) from exc
        raise InvalidAttachmentError("Invalid or corrupt PDF file") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:
            raise InvalidAttachmentError(
                "Password-protected or encrypted PDFs are not supported"
            ) from exc
        # pypdf returns 0 when the password fails; non-zero / True when unlocked.
        if unlocked in (0, False):
            raise InvalidAttachmentError(
                "Password-protected or encrypted PDFs are not supported"
            )

    try:
        # Touch page count early so encrypted/corrupt structures fail before extraction.
        _ = len(reader.pages)
    except FileNotDecryptedError as exc:
        raise InvalidAttachmentError(
            "Password-protected or encrypted PDFs are not supported"
        ) from exc
    except Exception as exc:
        raise InvalidAttachmentError("Invalid or corrupt PDF file") from exc

    return reader


def _pdf_reader_to_text(reader) -> str:
    parts: list[str] = []
    truncated = False
    page_count = len(reader.pages)
    pages_to_scan = min(page_count, _PDF_MAX_PAGES)
    if page_count > _PDF_MAX_PAGES:
        truncated = True

    for index in range(pages_to_scan):
        try:
            raw = reader.pages[index].extract_text() or ""
        except Exception as exc:
            message = str(exc).lower()
            if "password" in message or "encrypt" in message or "decrypt" in message:
                raise InvalidAttachmentError(
                    "Password-protected or encrypted PDFs are not supported"
                ) from exc
            raise InvalidAttachmentError("Invalid or corrupt PDF file") from exc

        page_text = _BLANK_LINE_RE.sub("\n\n", raw).strip()
        if not page_text:
            continue
        parts.append(f"[Page {index + 1}]\n{page_text}")
        if len("\n\n".join(parts)) >= ATTACHMENT_TEXT_EXCERPT_MAX:
            truncated = True
            break

    text = "\n\n".join(parts).strip()
    if truncated and text:
        budget = ATTACHMENT_TEXT_EXCERPT_MAX - len(_TRUNCATION_MARKER)
        if budget < 0:
            return _TRUNCATION_MARKER.strip()
        text = text[:budget].rstrip() + _TRUNCATION_MARKER
    return text
