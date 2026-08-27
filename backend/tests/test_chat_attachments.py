"""Chat attachment upload and turn-linking tests."""

from __future__ import annotations

from io import BytesIO

import pytest
from httpx import ASGITransport, AsyncClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.dependencies import AuthContext, get_auth_context
from app.core.exceptions import (
    ConflictError,
    InvalidAudioError,
    NotFoundError,
    SilentAudioError,
    TranscriptionBusyError,
    TranscriptionTimeoutError,
)
from app.db.models import Chat, ChatAttachment, Turn
from app.db.session import get_db
from app.main import create_app
from app.schemas.api import TurnCreateRequest
from app.services.chat_service import chat_service
from tests.conftest import create_model_set, create_other_auth


async def _create_chat(db: AsyncSession, auth: AuthContext) -> Chat:
    chat = Chat(org_id=auth.org_id, created_by=auth.user.id, title="Upload chat")
    db.add(chat)
    await db.flush()
    return chat


async def _client_for(db: AsyncSession, auth: AuthContext) -> AsyncClient:
    app = create_app()

    async def override_db():
        yield db

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_auth_context] = lambda: auth
    return AsyncClient(transport=ASGITransport(app=app), base_url="http://test")


async def _upload(
    client: AsyncClient,
    chat_id: str,
    *,
    filename: str = "notes.txt",
    content: bytes = b"hello multimind",
    content_type: str = "text/plain",
) -> object:
    return await client.post(
        f"/api/v1/chats/{chat_id}/attachments",
        files={"file": (filename, BytesIO(content), content_type)},
    )


@pytest.mark.asyncio
async def test_text_upload_persists_metadata_and_file(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    monkeypatch.setattr(get_settings(), "chat_attachment_max_bytes", 1024)

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"hello world")

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "notes.txt"
    assert body["excerpt_status"] == "ready"
    assert body["text_excerpt"] == "hello world"
    assert body["size_bytes"] == len(b"hello world")

    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == body["id"]))
    ).scalar_one()
    assert row.org_id == auth.org_id
    assert row.chat_id == chat.id
    assert row.uploaded_by_user_id == auth.user.id
    assert row.turn_id is None
    assert (attach_dir / row.relative_path).is_file()
    assert (attach_dir / row.relative_path).read_bytes() == b"hello world"


@pytest.mark.asyncio
async def test_unsupported_extension_returns_415(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="malware.exe",
            content=b"MZ",
            content_type="application/octet-stream",
        )
    assert response.status_code == 415
    assert response.json()["error"] == "UNSUPPORTED_ATTACHMENT_TYPE"


@pytest.mark.asyncio
async def test_unsupported_mime_returns_415(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="notes.txt",
            content=b"hello",
            content_type="application/pdf",
        )
    assert response.status_code == 415
    assert response.json()["error"] == "UNSUPPORTED_ATTACHMENT_TYPE"


@pytest.mark.asyncio
async def test_empty_file_rejected(db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"")
    assert response.status_code == 422
    assert response.json()["error"] == "INVALID_ATTACHMENT"


@pytest.mark.asyncio
async def test_oversized_file_rejected(db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    monkeypatch.setattr(get_settings(), "chat_attachment_max_bytes", 8)
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"0123456789")
    assert response.status_code == 413
    assert response.json()["error"] == "ATTACHMENT_TOO_LARGE"


@pytest.mark.asyncio
async def test_pending_list_is_org_and_chat_scoped(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    other_chat = await _create_chat(db, auth)
    other_auth = await create_other_auth(db)
    other_org_chat = Chat(
        org_id=other_auth.org_id, created_by=other_auth.user.id, title="Other org"
    )
    db.add(other_org_chat)
    await db.flush()

    async with await _client_for(db, auth) as client:
        first = await _upload(client, chat.id, filename="a.txt", content=b"one")
        second = await _upload(client, chat.id, filename="b.txt", content=b"two")
        await _upload(client, other_chat.id, filename="c.txt", content=b"three")
        listed = await client.get(f"/api/v1/chats/{chat.id}/attachments")

    async with await _client_for(db, other_auth) as other_client:
        await _upload(other_client, other_org_chat.id, filename="d.txt", content=b"four")
        other_listed = await other_client.get(f"/api/v1/chats/{chat.id}/attachments")

    assert first.status_code == 201
    assert second.status_code == 201
    assert listed.status_code == 200
    ids = [item["id"] for item in listed.json()["items"]]
    assert set(ids) == {first.json()["id"], second.json()["id"]}
    assert len(ids) == 2
    assert other_listed.status_code == 404


@pytest.mark.asyncio
async def test_turn_creation_links_attachments_and_builds_context(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)

    async with await _client_for(db, auth) as client:
        first = await _upload(client, chat.id, filename="alpha.txt", content=b"alpha content")
        second = await _upload(client, chat.id, filename="beta.md", content=b"beta content")

    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="Summarize the files",
            model_set_id="research-set",
            custom_instructions="Reference prior chat context.",
            attachment_ids=[second.json()["id"], first.json()["id"]],
        ),
    )
    await db.commit()

    assert {
        (attachment.id, attachment.filename, attachment.content_type)
        for attachment in turn.attachments
    } == {
        (first.json()["id"], "alpha.txt", "text/plain"),
        (second.json()["id"], "beta.md", "text/plain"),
    }

    reloaded_turns = await chat_service.list_turns(db, auth, chat.id)
    reloaded = next(item for item in reloaded_turns if item.id == turn.id)
    assert {
        (attachment.id, attachment.filename, attachment.content_type)
        for attachment in reloaded.attachments
    } == {
        (first.json()["id"], "alpha.txt", "text/plain"),
        (second.json()["id"], "beta.md", "text/plain"),
    }

    stored = (await db.execute(select(Turn).where(Turn.id == turn.id))).scalar_one()
    assert stored.custom_instructions is not None
    assert "Reference prior chat context." in stored.custom_instructions
    beta_pos = stored.custom_instructions.index("beta.md")
    alpha_pos = stored.custom_instructions.index("alpha.txt")
    assert beta_pos < alpha_pos
    assert "beta content" in stored.custom_instructions
    assert "alpha content" in stored.custom_instructions

    attachment_ids = [first.json()["id"], second.json()["id"]]
    rows = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id.in_(attachment_ids)))
    ).scalars().all()
    assert {row.turn_id for row in rows} == {turn.id}


@pytest.mark.asyncio
async def test_cross_org_attachment_id_rejected(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    other_auth = await create_other_auth(db)
    await create_model_set(db, other_auth, slug="research-set")

    chat = await _create_chat(db, auth)
    other_chat = Chat(org_id=other_auth.org_id, created_by=other_auth.user.id, title="Other")
    db.add(other_chat)
    await db.flush()

    async with await _client_for(db, other_auth) as other_client:
        uploaded = await _upload(other_client, other_chat.id, content=b"secret")

    with pytest.raises(NotFoundError):
        await chat_service.start_turn(
            db,
            auth,
            chat.id,
            TurnCreateRequest(
                user_message="Leak?",
                model_set_id="research-set",
                attachment_ids=[uploaded.json()["id"]],
            ),
        )


@pytest.mark.asyncio
async def test_attachment_from_different_chat_rejected(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    other_chat = await _create_chat(db, auth)

    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, other_chat.id, content=b"wrong chat")

    with pytest.raises(NotFoundError):
        await chat_service.start_turn(
            db,
            auth,
            chat.id,
            TurnCreateRequest(
                user_message="Use other chat file",
                model_set_id="research-set",
                attachment_ids=[uploaded.json()["id"]],
            ),
        )


@pytest.mark.asyncio
async def test_already_linked_attachment_rejected(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)

    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"once")

    await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="First",
            model_set_id="research-set",
            attachment_ids=[uploaded.json()["id"]],
        ),
    )
    await db.flush()

    with pytest.raises(ConflictError):
        await chat_service.start_turn(
            db,
            auth,
            chat.id,
            TurnCreateRequest(
                user_message="Second",
                model_set_id="research-set",
                attachment_ids=[uploaded.json()["id"]],
            ),
        )


@pytest.mark.asyncio
async def test_multiple_attachments_preserve_request_order(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)

    async with await _client_for(db, auth) as client:
        a = await _upload(client, chat.id, filename="a.txt", content=b"A")
        b = await _upload(client, chat.id, filename="b.txt", content=b"B")
        c = await _upload(client, chat.id, filename="c.txt", content=b"C")

    ordered_ids = [c.json()["id"], a.json()["id"], b.json()["id"]]
    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="Order check",
            model_set_id="research-set",
            attachment_ids=ordered_ids,
        ),
    )
    stored = (await db.execute(select(Turn).where(Turn.id == turn.id))).scalar_one()
    assert stored.custom_instructions is not None
    positions = [stored.custom_instructions.index(name) for name in ("c.txt", "a.txt", "b.txt")]
    assert positions == sorted(positions)


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _docx_bytes(
    *,
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
) -> bytes:
    from docx import Document

    document = Document()
    for text in paragraphs or []:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = cell
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes(
    *,
    sheets: dict[str, list[list[object]]] | None = None,
    formulas: dict[str, str] | None = None,
) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.remove(workbook.active)
    for name, rows in (sheets or {"Sheet1": [["value"]]}).items():
        worksheet = workbook.create_sheet(name)
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, value in enumerate(row, start=1):
                worksheet.cell(row=r_idx, column=c_idx, value=value)
    if formulas:
        sheet_name, formula = next(iter(formulas.items()))
        workbook[sheet_name]["A1"] = formula
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_docx_upload_extracts_paragraphs_and_tables(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    content = _docx_bytes(
        paragraphs=["Hello paragraph", ""],
        table_rows=[["Name", "Score"], ["Ada", "10"]],
    )
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="brief.docx",
            content=content,
            content_type=DOCX_MIME,
        )
    assert response.status_code == 201
    body = response.json()
    assert body["excerpt_status"] == "ready"
    assert "Hello paragraph" in body["text_excerpt"]
    assert "Name\tScore" in body["text_excerpt"]
    assert "Ada\t10" in body["text_excerpt"]


@pytest.mark.asyncio
async def test_empty_docx_has_empty_excerpt_status(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="empty.docx",
            content=_docx_bytes(paragraphs=[]),
            content_type=DOCX_MIME,
        )
    assert response.status_code == 201
    assert response.json()["excerpt_status"] == "empty"
    assert response.json()["text_excerpt"] is None


@pytest.mark.asyncio
async def test_corrupt_docx_returns_422_without_row_or_file(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="bad.docx",
            content=b"not-a-docx",
            content_type=DOCX_MIME,
        )
    assert response.status_code == 422
    assert response.json()["error"] == "INVALID_ATTACHMENT"
    assert (await db.execute(select(ChatAttachment))).scalars().all() == []
    assert list(attach_dir.rglob("*")) == [] or not any(attach_dir.rglob("*.docx"))


@pytest.mark.asyncio
async def test_xlsx_upload_extracts_sheets_and_values(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    content = _xlsx_bytes(sheets={"Sales": [["Region", "Total"], ["EU", 12]]})
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="sales.xlsx",
            content=content,
            content_type=XLSX_MIME,
        )
    assert response.status_code == 201
    body = response.json()
    assert body["excerpt_status"] == "ready"
    assert "Worksheet: Sales" in body["text_excerpt"]
    assert "Region\tTotal" in body["text_excerpt"]
    assert "EU\t12" in body["text_excerpt"]


@pytest.mark.asyncio
async def test_xlsx_formulas_are_not_executed(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    content = _xlsx_bytes(sheets={"Calc": []}, formulas={"Calc": "=1+1"})
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="calc.xlsx",
            content=content,
            content_type=XLSX_MIME,
        )
    assert response.status_code == 201
    excerpt = response.json()["text_excerpt"] or ""
    assert "=1+1" not in excerpt
    assert "2" not in excerpt.split()


@pytest.mark.asyncio
async def test_xlsx_respects_sheet_row_column_limits(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    sheets = {f"S{i}": [[f"v{i}-{c}" for c in range(35)] for _ in range(5)] for i in range(7)}
    content = _xlsx_bytes(sheets=sheets)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="wide.xlsx",
            content=content,
            content_type=XLSX_MIME,
        )
    assert response.status_code == 201
    excerpt = response.json()["text_excerpt"]
    assert "[Content truncated]" in excerpt
    assert "Worksheet: S0" in excerpt
    assert "Worksheet: S5" not in excerpt
    assert "v0-29" in excerpt
    assert "v0-30" not in excerpt


@pytest.mark.asyncio
async def test_corrupt_xlsx_returns_422_without_row_or_file(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="bad.xlsx",
            content=b"not-xlsx",
            content_type=XLSX_MIME,
        )
    assert response.status_code == 422
    assert response.json()["error"] == "INVALID_ATTACHMENT"
    assert (await db.execute(select(ChatAttachment))).scalars().all() == []
    assert not any(attach_dir.rglob("*.xlsx"))


@pytest.mark.asyncio
async def test_legacy_doc_and_xls_rejected(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        doc = await _upload(
            client, chat.id, filename="old.doc", content=b"x", content_type="application/msword"
        )
        xls = await _upload(
            client,
            chat.id,
            filename="old.xls",
            content=b"x",
            content_type="application/vnd.ms-excel",
        )
    assert doc.status_code == 415
    assert xls.status_code == 415


@pytest.mark.asyncio
async def test_octet_stream_accepted_only_for_valid_office_files(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        ok_docx = await _upload(
            client,
            chat.id,
            filename="ok.docx",
            content=_docx_bytes(paragraphs=["streamed"]),
            content_type="application/octet-stream",
        )
        bad_docx = await _upload(
            client,
            chat.id,
            filename="bad.docx",
            content=b"nope",
            content_type="application/octet-stream",
        )
        ok_xlsx = await _upload(
            client,
            chat.id,
            filename="ok.xlsx",
            content=_xlsx_bytes(sheets={"A": [["1"]]}),
            content_type="application/octet-stream",
        )
    assert ok_docx.status_code == 201
    assert "streamed" in ok_docx.json()["text_excerpt"]
    assert bad_docx.status_code == 422
    assert ok_xlsx.status_code == 201


@pytest.mark.asyncio
async def test_office_excerpt_reaches_turn_instructions(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        docx = await _upload(
            client,
            chat.id,
            filename="memo.docx",
            content=_docx_bytes(paragraphs=["Council memo body"]),
            content_type=DOCX_MIME,
        )
        xlsx = await _upload(
            client,
            chat.id,
            filename="grid.xlsx",
            content=_xlsx_bytes(sheets={"Grid": [["k", "v"], ["a", "1"]]}),
            content_type=XLSX_MIME,
        )
    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="Use both files",
            model_set_id="research-set",
            attachment_ids=[docx.json()["id"], xlsx.json()["id"]],
        ),
    )
    stored = (await db.execute(select(Turn).where(Turn.id == turn.id))).scalar_one()
    assert "Council memo body" in (stored.custom_instructions or "")
    assert "Worksheet: Grid" in (stored.custom_instructions or "")
    assert "a\t1" in (stored.custom_instructions or "")


@pytest.mark.asyncio
async def test_delete_pending_attachment_removes_row_and_file(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"delete me")
        attachment_id = uploaded.json()["id"]
        row = (
            await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
        ).scalar_one()
        stored_path = attach_dir / row.relative_path
        assert stored_path.is_file()

        response = await client.delete(
            f"/api/v1/chats/{chat.id}/attachments/{attachment_id}"
        )

    assert response.status_code == 200
    assert (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
    ).scalar_one_or_none() is None
    assert not stored_path.exists()


@pytest.mark.asyncio
async def test_delete_attachment_rejects_other_org(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    other_auth = await create_other_auth(db)
    other_chat = Chat(org_id=other_auth.org_id, created_by=other_auth.user.id, title="Other")
    db.add(other_chat)
    await db.flush()

    async with await _client_for(db, other_auth) as other_client:
        uploaded = await _upload(other_client, other_chat.id, content=b"secret")

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await client.delete(
            f"/api/v1/chats/{chat.id}/attachments/{uploaded.json()['id']}"
        )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_delete_attachment_rejects_other_chat(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    other_chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"mine")
        response = await client.delete(
            f"/api/v1/chats/{other_chat.id}/attachments/{uploaded.json()['id']}"
        )
    assert response.status_code == 404
    assert (
        await db.execute(
            select(ChatAttachment).where(ChatAttachment.id == uploaded.json()["id"])
        )
    ).scalar_one_or_none() is not None


@pytest.mark.asyncio
async def test_delete_linked_attachment_returns_conflict(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"linked")
        attachment_id = uploaded.json()["id"]

    await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="use file",
            model_set_id="research-set",
            attachment_ids=[attachment_id],
        ),
    )
    await db.commit()

    async with await _client_for(db, auth) as client:
        response = await client.delete(
            f"/api/v1/chats/{chat.id}/attachments/{attachment_id}"
        )
    assert response.status_code == 409
    assert response.json()["error"] == "CONFLICT"


@pytest.mark.asyncio
async def test_delete_attachment_missing_disk_file_still_deletes_row(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"gone file")
        attachment_id = uploaded.json()["id"]
        row = (
            await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
        ).scalar_one()
        (attach_dir / row.relative_path).unlink()

        response = await client.delete(
            f"/api/v1/chats/{chat.id}/attachments/{attachment_id}"
        )

    assert response.status_code == 200
    assert (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
    ).scalar_one_or_none() is None


@pytest.mark.asyncio
async def test_delete_attachment_never_deletes_out_of_root_path(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    attach_dir.mkdir()
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    outside = tmp_path / "outside-secret.txt"
    outside.write_text("do not delete", encoding="utf-8")

    chat = await _create_chat(db, auth)
    row = ChatAttachment(
        org_id=auth.org_id,
        chat_id=chat.id,
        uploaded_by_user_id=auth.user.id,
        turn_id=None,
        filename="escape.txt",
        stored_name="escape.txt",
        content_type="text/plain",
        size_bytes=4,
        relative_path="../outside-secret.txt",
        text_excerpt="x",
        excerpt_status="ready",
    )
    db.add(row)
    await db.flush()
    attachment_id = row.id

    async with await _client_for(db, auth) as client:
        response = await client.delete(
            f"/api/v1/chats/{chat.id}/attachments/{attachment_id}"
        )

    assert response.status_code == 200
    assert outside.read_text(encoding="utf-8") == "do not delete"
    assert (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
    ).scalar_one_or_none() is None


@pytest.mark.asyncio
async def test_delete_chat_removes_attachment_files(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        first = await _upload(client, chat.id, filename="a.txt", content=b"one")
        second = await _upload(client, chat.id, filename="b.txt", content=b"two")

    rows = (
        await db.execute(
            select(ChatAttachment).where(
                ChatAttachment.id.in_([first.json()["id"], second.json()["id"]])
            )
        )
    ).scalars().all()
    paths = [attach_dir / row.relative_path for row in rows]
    assert all(path.is_file() for path in paths)

    await chat_service.delete_chat(db, auth, chat.id)

    assert all(not path.exists() for path in paths)
    remaining = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.chat_id == chat.id))
    ).scalars().all()
    assert remaining == []


@pytest.mark.asyncio
async def test_delete_chat_succeeds_when_one_attachment_file_missing(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        first = await _upload(client, chat.id, filename="keep.txt", content=b"one")
        second = await _upload(client, chat.id, filename="gone.txt", content=b"two")

    first_row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == first.json()["id"]))
    ).scalar_one()
    second_row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == second.json()["id"]))
    ).scalar_one()
    first_path = attach_dir / first_row.relative_path
    second_path = attach_dir / second_row.relative_path
    second_path.unlink()

    await chat_service.delete_chat(db, auth, chat.id)

    assert not first_path.exists()
    assert (
        await db.execute(select(Chat).where(Chat.id == chat.id))
    ).scalar_one_or_none() is None


@pytest.mark.asyncio
async def test_soft_deleted_turn_keeps_linked_attachment_out_of_pending(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"linked forever")
        attachment_id = uploaded.json()["id"]

    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="use file",
            model_set_id="research-set",
            attachment_ids=[attachment_id],
        ),
    )
    await db.commit()

    await chat_service.delete_turn(db, auth, chat.id, turn.id)

    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
    ).scalar_one()
    assert row.turn_id == turn.id

    async with await _client_for(db, auth) as client:
        listed = await client.get(f"/api/v1/chats/{chat.id}/attachments")
    assert listed.status_code == 200
    assert listed.json()["items"] == []


@pytest.mark.asyncio
async def test_absolute_attachment_root_works_like_docker_path(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = (tmp_path / "app" / "data" / "chat_attachments").resolve()
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"docker root")

    assert response.status_code == 201
    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == response.json()["id"]))
    ).scalar_one()
    stored = attach_dir / row.relative_path
    assert stored.is_file()
    assert stored.read_bytes() == b"docker root"


@pytest.mark.asyncio
async def test_upload_just_below_size_limit_succeeds(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    monkeypatch.setattr(get_settings(), "chat_attachment_max_bytes", 64)
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"x" * 64)
    assert response.status_code == 201
    assert response.json()["size_bytes"] == 64


@pytest.mark.asyncio
async def test_upload_just_above_size_limit_returns_413_and_clears_temp(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    monkeypatch.setattr(get_settings(), "chat_attachment_max_bytes", 64)
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"x" * 65)
    assert response.status_code == 413
    assert response.json()["error"] == "ATTACHMENT_TOO_LARGE"
    tmp_dir = attach_dir / ".tmp"
    leftovers = list(tmp_dir.glob("*.upload*")) if tmp_dir.exists() else []
    assert leftovers == []
    assert list(attach_dir.rglob("*.txt")) == []


@pytest.mark.asyncio
async def test_extraction_failure_removes_temp_and_skips_db_row(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="bad.docx",
            content=b"not-a-docx",
            content_type=DOCX_MIME,
        )
    assert response.status_code == 422
    body = response.json()
    assert body["error"] == "INVALID_ATTACHMENT"
    assert ":\\" not in body["message"]
    assert "/tmp/" not in body["message"]
    assert str(attach_dir) not in body["message"]
    assert (await db.execute(select(ChatAttachment))).scalars().all() == []
    tmp_dir = attach_dir / ".tmp"
    leftovers = list(tmp_dir.glob("*.upload*")) if tmp_dir.exists() else []
    assert leftovers == []


@pytest.mark.asyncio
async def test_db_insert_failure_removes_final_file(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)

    original_flush = db.flush

    async def failing_flush(*args, **kwargs):
        for obj in list(db.new):
            if isinstance(obj, ChatAttachment):
                raise RuntimeError("simulated db failure")
        return await original_flush(*args, **kwargs)

    monkeypatch.setattr(db, "flush", failing_flush)

    async with await _client_for(db, auth) as client:
        with pytest.raises(RuntimeError, match="simulated db failure"):
            await _upload(client, chat.id, content=b"persist me")
    await db.rollback()
    assert list(attach_dir.rglob("*.txt")) == []
    assert (
        await db.execute(select(ChatAttachment))
    ).scalars().all() == []


@pytest.mark.asyncio
async def test_uppercase_office_extensions_accepted(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        docx = await _upload(
            client,
            chat.id,
            filename="MEMO.DOCX",
            content=_docx_bytes(paragraphs=["Upper DOCX"]),
            content_type=DOCX_MIME,
        )
        xlsx = await _upload(
            client,
            chat.id,
            filename="GRID.XLSX",
            content=_xlsx_bytes(sheets={"S": [["a", "1"]]}),
            content_type=XLSX_MIME,
        )
    assert docx.status_code == 201
    assert xlsx.status_code == 201
    assert "Upper DOCX" in docx.json()["text_excerpt"]
    assert "Worksheet: S" in xlsx.json()["text_excerpt"]


def test_null_byte_filename_rejected():
    from app.api.v1 import chats as chats_api
    from app.core.exceptions import InvalidAttachmentError

    with pytest.raises(InvalidAttachmentError):
        chats_api._validate_attachment_filename("evil\x00.txt")
    with pytest.raises(InvalidAttachmentError):
        chats_api._validate_attachment_filename("   ")


@pytest.mark.asyncio
async def test_whitespace_only_text_is_empty_excerpt(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=b"  \n\t  ")
    assert response.status_code == 201
    assert response.json()["excerpt_status"] == "empty"
    assert response.json()["text_excerpt"] is None


@pytest.mark.asyncio
async def test_large_text_attachment_retains_up_to_100k_characters(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    content = ("A" * 60_000) + ("B" * 40_000) + ("C" * 10_000)
    async with await _client_for(db, auth) as client:
        response = await _upload(client, chat.id, content=content.encode())

    assert response.status_code == 201
    excerpt = response.json()["text_excerpt"]
    assert len(excerpt) == 100_000
    assert excerpt == content[:100_000]
    assert len(excerpt) > 50_000
    assert "B" * 40_000 in excerpt
    assert "C" not in excerpt


@pytest.mark.asyncio
async def test_empty_attachment_does_not_inject_fake_content_block(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, filename="blank.txt", content=b"\n\n")
    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="Use blank",
            model_set_id="research-set",
            attachment_ids=[uploaded.json()["id"]],
        ),
    )
    stored = (await db.execute(select(Turn).where(Turn.id == turn.id))).scalar_one()
    instructions = stored.custom_instructions or ""
    assert "blank.txt" in instructions
    assert "(No readable content)" in instructions
    assert "```text" not in instructions


def test_attachment_context_budget_truncates_deterministically(monkeypatch):
    monkeypatch.setattr(get_settings(), "chat_attachment_context_max_chars", 220)
    from types import SimpleNamespace

    attachments = [
        SimpleNamespace(
            filename="first.txt",
            excerpt_status="ready",
            text_excerpt="ALPHA" * 80,
        ),
        SimpleNamespace(
            filename="second.txt",
            excerpt_status="ready",
            text_excerpt="BETA" * 80,
        ),
    ]
    text = chat_service._build_attachment_instructions(attachments)
    assert text is not None
    assert "first.txt" in text
    assert "second.txt" in text
    assert (
        "[Attachment context truncated]" in text
        or "[Content omitted due to attachment context budget]" in text
    )
    assert text.index("second.txt") > text.index("first.txt")


@pytest.mark.asyncio
async def test_duplicate_display_names_create_distinct_records(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        first = await _upload(client, chat.id, filename="same.txt", content=b"one")
        second = await _upload(client, chat.id, filename="same.txt", content=b"two")
    assert first.status_code == 201
    assert second.status_code == 201
    assert first.json()["id"] != second.json()["id"]
    rows = (
        await db.execute(
            select(ChatAttachment).where(
                ChatAttachment.id.in_([first.json()["id"], second.json()["id"]])
            )
        )
    ).scalars().all()
    paths = {attach_dir / row.relative_path for row in rows}
    assert len(paths) == 2
    assert all(path.is_file() for path in paths)


@pytest.mark.asyncio
async def test_conditional_link_prevents_double_claim(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(client, chat.id, content=b"once only")
    attachment_id = uploaded.json()["id"]

    first = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="First claim",
            model_set_id="research-set",
            attachment_ids=[attachment_id],
        ),
    )
    await db.flush()

    with pytest.raises(ConflictError):
        await chat_service.start_turn(
            db,
            auth,
            chat.id,
            TurnCreateRequest(
                user_message="Second claim",
                model_set_id="research-set",
                attachment_ids=[attachment_id],
            ),
        )

    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
    ).scalar_one()
    assert row.turn_id == first.id


PDF_MIME = "application/pdf"


def _pdf_with_pages(texts: list[str]) -> bytes:
    """Build a minimal multi-page PDF with one Helvetica text line per page."""
    objects: list[bytes] = []

    def add(obj: bytes) -> int:
        objects.append(obj)
        return len(objects)

    font_obj = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    page_objs: list[int] = []
    for text in texts:
        safe = (
            text.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
            .encode("latin-1", errors="replace")
        )
        stream = b"BT /F1 12 Tf 72 720 Td (" + safe + b") Tj ET"
        content_obj = add(
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )
        page_obj = add(
            b"<< /Type /Page /Parent 0 0 R /MediaBox [0 0 612 792] /Contents "
            + f"{content_obj} 0 R".encode()
            + b" /Resources << /Font << /F1 "
            + f"{font_obj} 0 R".encode()
            + b" >> >> >>"
        )
        page_objs.append(page_obj)

    kids = b"[" + b" ".join(f"{pid} 0 R".encode() for pid in page_objs) + b"]"
    pages_obj = add(b"<< /Type /Pages /Kids " + kids + f" /Count {len(page_objs)} >>".encode())
    for page_id in page_objs:
        objects[page_id - 1] = objects[page_id - 1].replace(
            b"/Parent 0 0 R",
            f"/Parent {pages_obj} 0 R".encode(),
        )
    catalog_obj = add(f"<< /Type /Catalog /Pages {pages_obj} 0 R >>".encode())

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out.extend(f"{i} 0 obj\n".encode())
        out.extend(obj)
        out.extend(b"\nendobj\n")
    xref_pos = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    out.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode())
    out.extend(f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_obj} 0 R >>\n".encode())
    out.extend(f"startxref\n{xref_pos}\n%%EOF\n".encode())
    return bytes(out)


def _blank_pdf() -> bytes:
    from io import BytesIO

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=300, height=144)
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _password_pdf() -> bytes:
    from io import BytesIO

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=300, height=144)
    writer.encrypt("secret-password")
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_valid_text_pdf_upload_extracts_page_text(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="notes.pdf",
            content=_pdf_with_pages(["Council alpha", "Council beta"]),
            content_type=PDF_MIME,
        )
    assert response.status_code == 201
    body = response.json()
    assert body["excerpt_status"] == "ready"
    assert "[Page 1]" in body["text_excerpt"]
    assert "Council alpha" in body["text_excerpt"]
    assert "[Page 2]" in body["text_excerpt"]
    assert "Council beta" in body["text_excerpt"]
    assert body["text_excerpt"].index("[Page 1]") < body["text_excerpt"].index("[Page 2]")


@pytest.mark.asyncio
async def test_pdf_extraction_truncates_at_excerpt_limit(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    long_line = "WORD" * 10_000
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="big.pdf",
            content=_pdf_with_pages([long_line, long_line, long_line]),
            content_type=PDF_MIME,
        )
    assert response.status_code == 201
    excerpt = response.json()["text_excerpt"]
    assert len(excerpt) <= 100_000
    assert len(excerpt) > 50_000
    assert "[Content truncated]" in excerpt


@pytest.mark.asyncio
async def test_pdf_page_limit_is_bounded(monkeypatch):
    monkeypatch.setattr(
        "app.services.chat_attachment_text._PDF_MAX_PAGES",
        2,
    )
    from app.services.chat_attachment_text import extract_attachment_text

    pages = [f"PageBody{i}" for i in range(1, 6)]
    excerpt, status = extract_attachment_text(_pdf_with_pages(pages), ".pdf")
    assert status == "ready"
    assert "[Page 1]" in excerpt
    assert "[Page 2]" in excerpt
    assert "[Page 3]" not in excerpt
    assert "[Content truncated]" in excerpt


@pytest.mark.asyncio
async def test_empty_pdf_produces_empty_excerpt_status(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="scan.pdf",
            content=_blank_pdf(),
            content_type=PDF_MIME,
        )
    assert response.status_code == 201
    assert response.json()["excerpt_status"] == "empty"
    assert response.json()["text_excerpt"] is None

    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="Use scan",
            model_set_id="research-set",
            attachment_ids=[response.json()["id"]],
        ),
    )
    stored = (await db.execute(select(Turn).where(Turn.id == turn.id))).scalar_one()
    instructions = stored.custom_instructions or ""
    assert "scan.pdf" in instructions
    assert "No readable text found" in instructions
    assert "```text" not in instructions


@pytest.mark.asyncio
async def test_corrupt_pdf_returns_422_and_cleans_temp(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="bad.pdf",
            content=b"%PDF-1.4 not-a-real-pdf",
            content_type=PDF_MIME,
        )
    assert response.status_code == 422
    body = response.json()
    assert body["error"] == "INVALID_ATTACHMENT"
    assert str(attach_dir) not in body["message"]
    assert ":\\" not in body["message"]
    leftovers = list((attach_dir / ".tmp").glob("*.upload*")) if (attach_dir / ".tmp").exists() else []
    assert leftovers == []
    assert list(attach_dir.rglob("*.pdf")) == []


@pytest.mark.asyncio
async def test_password_protected_pdf_returns_422(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="locked.pdf",
            content=_password_pdf(),
            content_type=PDF_MIME,
        )
    assert response.status_code == 422
    assert response.json()["error"] == "INVALID_ATTACHMENT"
    assert "Password-protected" in response.json()["message"]


@pytest.mark.asyncio
async def test_octet_stream_accepted_only_for_valid_pdf(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        ok = await _upload(
            client,
            chat.id,
            filename="ok.pdf",
            content=_pdf_with_pages(["octet pdf"]),
            content_type="application/octet-stream",
        )
        bad = await _upload(
            client,
            chat.id,
            filename="fake.pdf",
            content=b"not-pdf-bytes",
            content_type="application/octet-stream",
        )
    assert ok.status_code == 201
    assert "octet pdf" in ok.json()["text_excerpt"]
    assert bad.status_code == 422


@pytest.mark.asyncio
async def test_pdf_content_reaches_turn_instructions(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    await create_model_set(db, auth, slug="research-set")
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(
            client,
            chat.id,
            filename="brief.pdf",
            content=_pdf_with_pages(["Verdict brief body"]),
            content_type=PDF_MIME,
        )
    turn = await chat_service.start_turn(
        db,
        auth,
        chat.id,
        TurnCreateRequest(
            user_message="Use pdf",
            model_set_id="research-set",
            attachment_ids=[uploaded.json()["id"]],
        ),
    )
    stored = (await db.execute(select(Turn).where(Turn.id == turn.id))).scalar_one()
    assert "brief.pdf" in (stored.custom_instructions or "")
    assert "Verdict brief body" in (stored.custom_instructions or "")
    assert "[Page 1]" in (stored.custom_instructions or "")


def test_pdf_context_obeys_total_attachment_budget(monkeypatch):
    monkeypatch.setattr(get_settings(), "chat_attachment_context_max_chars", 240)
    from types import SimpleNamespace

    text = chat_service._build_attachment_instructions(
        [
            SimpleNamespace(
                filename="first.pdf",
                excerpt_status="ready",
                text_excerpt="[Page 1]\n" + ("ALPHA" * 80),
            ),
            SimpleNamespace(
                filename="second.pdf",
                excerpt_status="ready",
                text_excerpt="[Page 1]\n" + ("BETA" * 80),
            ),
        ]
    )
    assert text is not None
    assert "first.pdf" in text
    assert "second.pdf" in text
    assert (
        "[Attachment context truncated]" in text
        or "[Content omitted due to attachment context budget]" in text
    )


@pytest.mark.parametrize(
    ("sizes", "minimum_retained"),
    [
        ([50_000, 50_000], [50_000, 40_000]),
        ([70_000, 40_000], [70_000, 20_000]),
        ([25_000] * 4, [25_000, 25_000, 25_000, 15_000]),
        ([20_000] * 10, [20_000, 20_000, 20_000, 20_000, 10_000]),
    ],
)
def test_multiple_attachments_share_one_total_context_budget(
    monkeypatch, sizes, minimum_retained
):
    monkeypatch.setattr(get_settings(), "chat_attachment_context_max_chars", 100_000)
    from types import SimpleNamespace

    markers = "ABCDEFGHIJ"
    text = chat_service._build_attachment_instructions(
        [
            SimpleNamespace(
                filename=f"attachment-{index}.txt",
                excerpt_status="ready",
                text_excerpt=markers[index] * size,
            )
            for index, size in enumerate(sizes)
        ]
    )

    assert text is not None
    assert len(text) <= get_settings().chat_attachment_context_max_chars
    for index, minimum in enumerate(minimum_retained):
        assert text.count(markers[index]) >= minimum
        assert f"attachment-{index}.txt" in text
    filenames = [text.index(f"attachment-{index}.txt") for index in range(len(sizes))]
    assert filenames == sorted(filenames)
    assert (
        "[Attachment context truncated]" in text
        or "[Content omitted due to attachment context budget]" in text
    )


@pytest.mark.asyncio
async def test_pending_attachment_cap_rejects_eleventh_upload(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        for i in range(10):
            response = await _upload(
                client,
                chat.id,
                filename=f"f{i}.txt",
                content=f"body-{i}".encode(),
            )
            assert response.status_code == 201
        eleventh = await _upload(client, chat.id, filename="overflow.txt", content=b"nope")
    assert eleventh.status_code == 400
    assert eleventh.json()["error"] == "VALIDATION_ERROR"
    assert "at most 10" in eleventh.json()["message"]


# --- .webm audio attachment transcription ---------------------------------


def _fake_webm_bytes() -> bytes:
    # Opaque payload is fine: Whisper is mocked; we only need a non-empty file on disk.
    return b"FAKEWEBM\x00audio-payload-for-attachment-tests"


def test_webm_content_type_normalization_and_allowlist():
    """MIME parameter stripping + labels browsers/OS use for MediaRecorder .webm."""
    from app.api.v1.chats import (
        _normalize_media_type,
        _validate_attachment_content_type,
    )
    from app.core.exceptions import UnsupportedAttachmentTypeError

    assert _normalize_media_type("audio/webm;codecs=opus") == "audio/webm"
    assert _normalize_media_type("  AUDIO/WEBM ; codecs=opus ") == "audio/webm"

    assert (
        _validate_attachment_content_type("audio/webm", "clip.webm", ".webm") == "audio/webm"
    )
    assert (
        _validate_attachment_content_type("audio/webm;codecs=opus", "clip.webm", ".webm")
        == "audio/webm"
    )
    assert (
        _validate_attachment_content_type("  Audio/WebM ;codecs=opus ", "clip.webm", ".webm")
        == "audio/webm"
    )
    assert (
        _validate_attachment_content_type("video/webm", "clip.webm", ".webm") == "video/webm"
    )
    assert (
        _validate_attachment_content_type("video/webm;codecs=vp8,opus", "clip.webm", ".webm")
        == "video/webm"
    )
    assert (
        _validate_attachment_content_type("application/octet-stream", "clip.webm", ".webm")
        == "application/octet-stream"
    )
    # Missing Content-Type: do not use mimetypes.guess_type(.webm) → video/webm reject path.
    assert _validate_attachment_content_type(None, "clip.webm", ".webm") == "audio/webm"
    assert _validate_attachment_content_type("", "clip.webm", ".webm") == "audio/webm"

    for bad in ("audio/mpeg", "audio/wav", "audio/mp4", "image/png", "text/plain"):
        try:
            _validate_attachment_content_type(bad, "clip.webm", ".webm")
            raise AssertionError(f"expected reject for {bad}")
        except UnsupportedAttachmentTypeError:
            pass

    try:
        _validate_attachment_content_type("audio/mpeg", "song.mp3", ".mp3")
        raise AssertionError("mp3 must remain unsupported")
    except UnsupportedAttachmentTypeError:
        pass

    # Document MIME behavior unchanged.
    assert (
        _validate_attachment_content_type("application/pdf", "a.pdf", ".pdf") == "application/pdf"
    )
    assert (
        _validate_attachment_content_type(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "a.docx",
            ".docx",
        )
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@pytest.mark.asyncio
async def test_webm_upload_accepts_codecs_parameter_and_video_webm(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        with_codecs = await _upload(
            client,
            chat.id,
            filename="rec-codecs.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm;codecs=opus",
        )
        as_video = await _upload(
            client,
            chat.id,
            filename="rec-video-label.webm",
            content=_fake_webm_bytes() + b"2",
            content_type="video/webm",
        )
        missing_mime = await client.post(
            f"/api/v1/chats/{chat.id}/attachments",
            # Omit explicit Content-Type: clients often guess video/webm for .webm.
            files={"file": ("rec-empty-mime.webm", BytesIO(_fake_webm_bytes() + b"3"))},
        )

    assert with_codecs.status_code == 201, with_codecs.text
    assert with_codecs.json()["content_type"] == "audio/webm"
    assert as_video.status_code == 201, as_video.text
    assert as_video.json()["content_type"] == "video/webm"
    assert missing_mime.status_code == 201, missing_mime.text
    # httpx/Starlette typically guess video/webm from the .webm filename when MIME is omitted.
    assert missing_mime.json()["content_type"] in {"audio/webm", "video/webm"}


@pytest.mark.asyncio
async def test_webm_upload_is_stored_without_transcription(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    import app.api.v1.chats as chats_api

    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))

    async def fake_transcribe(*_args, **_kwargs):
        raise AssertionError("upload must not invoke transcription")

    monkeypatch.setattr(chats_api.transcription_service, "transcribe", fake_transcribe)

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="clip.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "clip.webm"
    assert body["content_type"] == "audio/webm"
    assert body["excerpt_status"] == "failed"
    assert body["text_excerpt"] is None
    assert body["size_bytes"] == len(_fake_webm_bytes())
    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == body["id"]))
    ).scalar_one()
    stored = attach_dir / row.relative_path
    assert stored.is_file()
    assert stored.suffix == ".webm"
    assert stored.read_bytes() == _fake_webm_bytes()
    assert not (attach_dir / ".tmp").exists() or not any((attach_dir / ".tmp").iterdir())


@pytest.mark.asyncio
async def test_mp3_attachment_still_rejected(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(tmp_path / "attachments"))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="song.mp3",
            content=b"ID3fake",
            content_type="audio/mpeg",
        )
    assert response.status_code == 415
    assert response.json()["error"] == "UNSUPPORTED_ATTACHMENT_TYPE"


@pytest.mark.asyncio
async def test_webm_upload_ignores_transcription_unavailability(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    import app.api.v1.chats as chats_api

    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))

    async def fake_transcribe(*_args, **_kwargs):
        raise AssertionError("upload must not invoke transcription")

    monkeypatch.setattr(chats_api.transcription_service, "transcribe", fake_transcribe)

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="quiet.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )

    assert response.status_code == 201
    body = response.json()
    assert body["excerpt_status"] == "failed"
    assert body["text_excerpt"] is None
    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == body["id"]))
    ).scalar_one()
    assert (attach_dir / row.relative_path).is_file()


@pytest.mark.asyncio
async def test_webm_upload_does_not_call_busy_transcription_service(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    import app.api.v1.chats as chats_api

    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))

    async def fake_transcribe(*_args, **_kwargs):
        raise AssertionError("upload must not invoke transcription")

    monkeypatch.setattr(chats_api.transcription_service, "transcribe", fake_transcribe)

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="busy.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )

    assert response.status_code == 201
    row = (await db.execute(select(ChatAttachment))).scalar_one()
    assert (attach_dir / row.relative_path).read_bytes() == _fake_webm_bytes()
    tmp_dir = attach_dir / ".tmp"
    if tmp_dir.exists():
        assert list(tmp_dir.iterdir()) == []
    assert [p for p in attach_dir.rglob("*") if p.is_file()] == [attach_dir / row.relative_path]


@pytest.mark.asyncio
async def test_webm_upload_does_not_call_failing_transcription_service(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    import app.api.v1.chats as chats_api

    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))

    async def fake_transcribe(*_args, **_kwargs):
        raise AssertionError("upload must not invoke transcription")

    monkeypatch.setattr(chats_api.transcription_service, "transcribe", fake_transcribe)

    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        response = await _upload(
            client,
            chat.id,
            filename="bad.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )

    assert response.status_code == 201
    row = (await db.execute(select(ChatAttachment))).scalar_one()
    stored = attach_dir / row.relative_path
    assert stored.read_bytes() == _fake_webm_bytes()


@pytest.mark.asyncio
async def test_stored_webm_transcription_returns_full_response_without_mutation(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    from pathlib import Path

    import app.api.v1.chats as chats_api
    from app.services.transcription_service import TranscriptionResult

    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(
            client,
            chat.id,
            filename="saved.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )
        attachment_id = uploaded.json()["id"]
        called: dict[str, object] = {}

        def fake_inspect(path: Path):
            called["inspected_path"] = path
            return 8.5

        async def fake_transcribe(path: Path, *, language: str | None = None):
            called["path"] = path
            called["language"] = language
            return TranscriptionResult(
                text="  full editable transcript  ",
                language="en",
                language_probability=0.98,
                duration_seconds=8.5,
                processing_seconds=1.2,
            )

        monkeypatch.setattr(chats_api, "inspect_audio_duration", fake_inspect)
        monkeypatch.setattr(chats_api.transcription_service, "transcribe_nowait", fake_transcribe)
        response = await client.post(
            f"/api/v1/chats/{chat.id}/attachments/{attachment_id}/transcription",
            json={"language": "en"},
        )

    assert response.status_code == 200, response.text
    assert response.json() == {
        "text": "full editable transcript",
        "language": "en",
        "language_probability": 0.98,
        "duration_seconds": 8.5,
        "processing_seconds": 1.2,
    }
    row = (
        await db.execute(select(ChatAttachment).where(ChatAttachment.id == attachment_id))
    ).scalar_one()
    assert called["path"] == attach_dir / row.relative_path
    assert called["inspected_path"] == called["path"]
    assert called["language"] == "en"
    assert row.text_excerpt is None
    assert row.excerpt_status == "failed"
    assert (attach_dir / row.relative_path).read_bytes() == _fake_webm_bytes()
    assert (await db.execute(select(Turn))).scalars().all() == []


@pytest.mark.asyncio
async def test_attachment_transcription_rejects_wrong_chat_cross_org_and_non_webm(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    other_chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        audio = await _upload(
            client,
            chat.id,
            filename="saved.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )
        document = await _upload(client, chat.id)
        wrong_chat = await client.post(
            f"/api/v1/chats/{other_chat.id}/attachments/{audio.json()['id']}/transcription",
            json={"language": "en"},
        )
        non_audio = await client.post(
            f"/api/v1/chats/{chat.id}/attachments/{document.json()['id']}/transcription",
            json={"language": "en"},
        )

    other_auth = await create_other_auth(db)
    async with await _client_for(db, other_auth) as other_client:
        cross_org = await other_client.post(
            f"/api/v1/chats/{chat.id}/attachments/{audio.json()['id']}/transcription",
            json={"language": "en"},
        )

    assert wrong_chat.status_code == 404
    assert cross_org.status_code == 404
    assert non_audio.status_code == 422
    assert (attach_dir / (await db.get(ChatAttachment, audio.json()["id"])).relative_path).is_file()


@pytest.mark.asyncio
async def test_attachment_transcription_missing_file_preserves_row(
    db: AsyncSession, auth: AuthContext, tmp_path, monkeypatch
):
    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(
            client,
            chat.id,
            filename="missing.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )
        row = await db.get(ChatAttachment, uploaded.json()["id"])
        (attach_dir / row.relative_path).unlink()
        response = await client.post(
            f"/api/v1/chats/{chat.id}/attachments/{row.id}/transcription",
            json={"language": "en"},
        )

    assert response.status_code == 422
    assert response.json()["error"] == "INVALID_ATTACHMENT"
    assert await db.get(ChatAttachment, row.id) is not None


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("error_factory", "status_code", "error_code"),
    [
        (TranscriptionBusyError, 429, "TRANSCRIPTION_BUSY"),
        (TranscriptionTimeoutError, 504, "TRANSCRIPTION_TIMEOUT"),
        (SilentAudioError, 422, "SILENT_AUDIO"),
        (InvalidAudioError, 422, "INVALID_AUDIO"),
    ],
)
async def test_attachment_transcription_failure_preserves_saved_audio(
    db: AsyncSession,
    auth: AuthContext,
    tmp_path,
    monkeypatch,
    error_factory,
    status_code,
    error_code,
):
    import app.api.v1.chats as chats_api

    attach_dir = tmp_path / "attachments"
    monkeypatch.setattr(get_settings(), "chat_attachment_dir", str(attach_dir))
    monkeypatch.setattr(chats_api, "inspect_audio_duration", lambda _path: 2.0)

    async def fail_transcription(*_args, **_kwargs):
        raise error_factory()

    monkeypatch.setattr(
        chats_api.transcription_service,
        "transcribe_nowait",
        fail_transcription,
    )
    chat = await _create_chat(db, auth)
    async with await _client_for(db, auth) as client:
        uploaded = await _upload(
            client,
            chat.id,
            filename="retry.webm",
            content=_fake_webm_bytes(),
            content_type="audio/webm",
        )
        row = await db.get(ChatAttachment, uploaded.json()["id"])
        response = await client.post(
            f"/api/v1/chats/{chat.id}/attachments/{row.id}/transcription",
            json={"language": "en"},
        )

    assert response.status_code == status_code
    assert response.json()["error"] == error_code
    assert await db.get(ChatAttachment, row.id) is not None
    assert (attach_dir / row.relative_path).read_bytes() == _fake_webm_bytes()


def test_audio_transcript_enters_attachment_instructions():
    from types import SimpleNamespace

    attachments = [
        SimpleNamespace(
            filename="clip.webm",
            excerpt_status="ready",
            text_excerpt="spoken transcript for the council",
        ),
        SimpleNamespace(
            filename="notes.txt",
            excerpt_status="ready",
            text_excerpt="document body",
        ),
    ]
    text = chat_service._build_attachment_instructions(attachments)
    assert text is not None
    assert "clip.webm" in text
    assert "spoken transcript for the council" in text
    assert "notes.txt" in text
    assert "document body" in text
    assert "```text" in text
